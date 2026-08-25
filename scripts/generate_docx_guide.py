import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_styled_document(output_path):
    doc = Document()

    # Set standard margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Color Palette
    PRIMARY_COLOR = RGBColor(30, 41, 59)      # Slate 800
    SECONDARY_COLOR = RGBColor(37, 99, 235)   # Blue 600
    TEXT_COLOR = RGBColor(51, 65, 85)         # Slate 700
    MUTED_COLOR = RGBColor(100, 116, 139)     # Slate 500
    BG_LIGHT_HEX = "F8FAFC"
    HEADER_BG_HEX = "1E293B"
    BORDER_HEX = "CBD5E1"
    ACCENT_BG_HEX = "EFF6FF"

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = TEXT_COLOR
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)

    def add_title(text, subtitle=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

        if subtitle:
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_after = Pt(18)
            run_sub = p_sub.add_run(subtitle)
            run_sub.font.name = 'Calibri'
            run_sub.font.size = Pt(13)
            run_sub.font.italic = True
            run_sub.font.color.rgb = SECONDARY_COLOR

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(17)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
            r_bold.font.color.rgb = PRIMARY_COLOR
        r_text = p.add_run(text)
        r_text.font.color.rgb = TEXT_COLOR

    def add_callout(text, title=None):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        # Shading & left border
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT_BG_HEX}"/>')
        tcPr.append(shd)
        
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="none"/>\n'
            f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/>\n'
            f'  <w:bottom w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if title:
            r_title = p.add_run(title + "\n")
            r_title.bold = True
            r_title.font.color.rgb = SECONDARY_COLOR
        r_body = p.add_run(text)
        r_body.font.color.rgb = TEXT_COLOR
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def style_table(table, col_widths, col_alignments=None):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            
            for j, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[j])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                tcPr = cell._tc.get_or_add_tcPr()
                
                # Padding
                mar = parse_xml(
                    f'<w:tcMar {nsdecls("w")}>\n'
                    f'  <w:top w:w="120" w:type="dxa"/>\n'
                    f'  <w:bottom w:w="120" w:type="dxa"/>\n'
                    f'  <w:left w:w="160" w:type="dxa"/>\n'
                    f'  <w:right w:w="160" w:type="dxa"/>\n'
                    f'</w:tcMar>'
                )
                tcPr.append(mar)

                # Border styling
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>\n'
                    f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_HEX}"/>\n'
                    f'  <w:left w:val="none"/>\n'
                    f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{BORDER_HEX}"/>\n'
                    f'  <w:right w:val="none"/>\n'
                    f'</w:tcBorders>'
                )
                tcPr.append(borders)

                # Header row styling
                if i == 0:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG_HEX}"/>')
                    tcPr.append(shd)
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        if col_alignments:
                            p.alignment = col_alignments[j]
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(9.5)
                else:
                    # Alternate zebra striping
                    if i % 2 == 1:
                        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{BG_LIGHT_HEX}"/>')
                        tcPr.append(shd)
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        if col_alignments:
                            p.alignment = col_alignments[j]
                        for run in p.runs:
                            run.font.size = Pt(9.5)

    # --- DOCUMENT GENERATION ---

    add_title("LayoutLoom: Pitch Deck & Architecture Guide", 
              "Comprehensive Technical Whitepaper & Presentation Blueprint for Layout-Preserving Document Translation")

    add_callout(
        "Executive Summary: Translation is only 10% of the challenge—document reconstruction is the remaining 90%. "
        "LayoutLoom is a state-of-the-art document reconstruction system that translates complex PDFs into foreign languages "
        "while ensuring output documents remain visually indistinguishable from originals except for the translated text.",
        "CORE VALUE PROPOSITION"
    )

    # SECTION 1
    add_h1("1. The Core Problem in Document Localization")
    doc.add_paragraph(
        "Translating structured documents like legal contracts, government gazettes, research papers, and scanned invoices "
        "is notoriously difficult because text dynamics change drastically across languages and scripts:"
    )
    add_bullet("Expansion & Contraction: ", "German text expands by ~32% relative to English, while Chinese contracts by ~38%.")
    add_bullet("Script Physics & Typography: ", "Devanagari (Hindi) requires 8% higher line spacing for complex conjunct ascenders and descenders. Arabic requires right-to-left flow, contextual cursive letter joining, and reversed column ordering while maintaining embedded LTR numeric strings.")
    add_bullet("Collateral Visual Damage: ", "Standard PDF translation tools overwrite vector stamps, official seals, and table borders, or clip text when it overflows its original container.")
    add_bullet("Lack of Measurable Proof: ", "Existing tools lack mathematical verification, often returning visually degraded or illegible documents without warning.")

    # SECTION 2
    add_h1("2. Slide-by-Slide Pitch Deck Blueprint")
    doc.add_paragraph("Use this 11-slide structure for an impactful competition, pitch, or stakeholder presentation:")

    slides = [
        ("Slide 1: Title & Hook",
         "LayoutLoom: Pixel-Accurate, Layout-Preserving Document Translation",
         "Side-by-side comparison of a Hindi Government Notice with official vector seal translated into English with identical alignment.",
         [("Hook: ", "Translating text is easy. Putting translated text of different length, script, and writing direction back into the exact original space without moving artwork is the real challenge."),
          ("Key Message: ", "Sub-millimeter typographic reconstruction with mathematical verification.")],
         "Open with high confidence. Show the visual proof immediately before introducing tech jargon."),

        ("Slide 2: The Industry Problem",
         "Why Standard PDF Translators Fail",
         "Visual breakdown of common failure modes: clipped overflow, ruined tables, displaced vector seals, and font tofu boxes.",
         [("Text Expansion: ", "30%+ growth causes text boxes to collide with adjacent elements."),
          ("Script Differences: ", "RTL Arabic and Devanagari ligatures break standard rendering engines."),
          ("Destructive Overwrites: ", "Naïve text replacement destroys underlying artwork, stamps, and watermarks.")],
         "Emphasize that unverified translations create legal and regulatory compliance risks in government and enterprise."),

        ("Slide 3: End-to-End System Architecture",
         "Modular, High-Performance Document Pipeline",
         "Architectural flowchart: RawDict Parsing → Classifier → Bidi Shaping → 7-Rung Fit Ladder → Masked SSIM Oracle.",
         [("In-Memory Layout Cache: ", "Indexed by SHA-256 to avoid redundant multi-pass parsing."),
          ("Immutable Copy-on-Write: ", "Every user edit generates an immutable version row without modifying original assets."),
          ("Zero-Dependency Offline Execution: ", "Full pipeline functions offline without external cloud API dependencies.")],
         "Highlight the clean separation between parsing, translation, layout fitting, and validation."),

        ("Slide 4: Text Extraction & Reading Order",
         "Intelligent Geometric Extraction",
         "Diagram showing span baseline detection and corridor gutter analysis on multi-column layouts.",
         [("RawDict Span Extraction: ", "Captures exact baseline coordinates, font families, point sizes, ascenders, descenders, and colors."),
          ("Multi-Stream Line Grouping: ", "Maintains multiple concurrent paragraph streams to handle locally multi-column bands cleanly."),
          ("Parallelism Gutter Checks: ", "Ensures true multi-column separation without mistaking table cell whitespace for columns.")],
         "Explain that baseline registration is what guarantees zero vertical drift."),

        ("Slide 5: Innovation 1 — The 7-Rung Typographic Fit Ladder",
         "Adaptive Geometry Without Clipping",
         "Visual diagram of the 7 rungs showing leading tightening, font scaling, and empty space expansion.",
         [("Rung 0: ", "Original font size and leading (0 penalty weight)."),
          ("Rungs 1-3: ", "Micro-concessions: tighten leading (floor 0.95x), tracking (-2% em), and font size reduction (floor 0.82x, 0.75x in tables)."),
          ("Rungs 4-5: ", "Spatial expansion: grow box downward into verified empty space (max +25%) or into outer margins."),
          ("Rung 6: ", "Honest overflow reporting: text is never clipped or hidden; exact deficits are logged.")],
         "This ladder mirrors the exact decision-making process of human typographers."),

        ("Slide 6: Innovation 2 — Font Resolution & Script Shaping",
         "Universal Script Handling & Optical Matching",
         "Glyph matrix showing Noto font families across Latin, Devanagari, Arabic, Japanese, and Chinese.",
         [("Optical X-Height Correction: ", "Normalizes replacement font sizes based on true physical x-height ratios [0.85, 1.15]."),
          ("Arabic Pre-Shaping & Bidi: ", "Applies Unicode Bidirectional Algorithm and ligature shaping before text box placement."),
          ("Devanagari Grapheme Clustering: ", "Prevents breaking conjunct characters across line boundaries via \\X regex splitting."),
          ("UAX #14 CJK Wrapping: ", "Calculates line breaks against advance widths while preventing orphaned punctuation.")],
         "Mention that all fonts are pre-vendored to eliminate tofu boxes."),

        ("Slide 7: Innovation 3 — Table Detection & Scanned OCR Inpainting",
         "Precision Tables & Scan Reconstruction",
         "Before/after comparison of a 200 DPI skewed invoice reconstructed with background inpainting.",
         [("3-Tier Table Classification: ", "Grid-based (interior rules), Rules-only (horizontal lines), and Alignment-only (unruled text clusters)."),
          ("Hough Line Deskewing: ", "Detects and corrects rotational skew at the raster's native resolution."),
          ("Per-Block Background Inpainting: ", "Samples a 3px outer ring around ink boundaries to match scan lighting without erasing grid lines."),
          ("Confidence Heatmap: ", "Flags OCR blocks under 75% confidence for interactive human review.")],
         "Show how inpainting preserves original scanned paper texture while placing crisp translated text."),

        ("Slide 8: Innovation 4 — Mathematical Validation & Masked SSIM",
         "Empirical Proof: Verified Without Hallucinations",
         "Formula callout with masked visual difference maps isolating vector artwork preservation.",
         [("Masked SSIM (>= 0.98): ", "Calculates Structural Similarity on all pixels outside text bboxes to prove artwork is untouched."),
          ("Adjustment Budget: ", "Fibonacci-weighted penalty score tracking typographic concessions."),
          ("Composite Layout Score: ", "Weighted formula combining SSIM, overflow rate, concession budget, and geometric integrity.")],
         "Stress that every score in the UI is clickable down to its exact mathematical derivation."),

        ("Slide 9: Empirical Benchmark Results",
         "Measured Performance Across Complex Real-World Documents",
         "Benchmark summary table showing 1.000 Masked SSIM across all 4 sample documents.",
         [("Govt Notice (Hindi): ", "Masked SSIM 1.000 | Layout Score: 89.1 | Overlaps: 0"),
          ("Research Paper (English/Hindi): ", "Masked SSIM 1.000 | Layout Score: 98.6 | Overlaps: 0"),
          ("Technical Report (German): ", "Masked SSIM 1.000 | Layout Score: 98.2 | Overlaps: 0"),
          ("Scanned Invoice (Scanned Scan): ", "Masked SSIM 1.000 | Layout Score: 92.6 | Overlaps: 0")],
         "These benchmarks prove the engine handles both pristine digital PDFs and noisy scans."),

        ("Slide 10: Interactive Comparison Studio (UI)",
         "Professional Review & Editing Environment",
         "Screenshot of the 3-pane comparison studio with synced zoom and issue inspector.",
         [("Triple-View Comparison: ", "Side-by-side, difference overlay, and raw text comparison modes."),
          ("Synchronized Pan & Zoom: ", "Pan on the source document instantly aligns the translated document canvas."),
          ("Live Issue Inspector: ", "Visual badges for font substitutions, OCR confidence, and overflows."),
          ("Surgical In-Place Edits: ", "Edit translated text blocks with instant single-page re-rendering.")],
         "Demonstrate the ease with which reviewers can audit and correct individual segments."),

        ("Slide 11: Business Impact & Roadmap",
         "Enterprise Viability & Extensibility",
         "Verticals diagram: Legal, Government, Academia, Localization.",
         [("Target Verticals: ", "Government records, cross-border legal contracts, patent filing, technical documentation."),
          ("Pluggable Adapters: ", "Extensible interfaces for Cloud OCR (AWS Textract, Google Vision) and AI Engines (OpenAI, Gemini, Ollama)."),
          ("Data Privacy: ", "Air-gapped deployment capability ensures zero data leakage.")],
         "Conclude with a strong pitch for enterprise adoption and privacy compliance.")
    ]

    for title, subtitle, visual, bullets, notes in slides:
        add_h2(title)
        p_sub = doc.add_paragraph()
        r_sub = p_sub.add_run(subtitle)
        r_sub.bold = True
        r_sub.font.color.rgb = SECONDARY_COLOR
        
        add_bullet("Suggested Visual: ", visual)
        for b_title, b_text in bullets:
            add_bullet(b_title, b_text)
        add_bullet("Speaker Note / Delivery Tip: ", notes)

    # SECTION 3
    add_h1("3. Comprehensive Technology Stack")
    doc.add_paragraph("LayoutLoom leverages a carefully selected modern stack designed for high throughput, memory safety, and deterministic rendering:")

    tech_table = doc.add_table(rows=11, cols=3)
    tech_data = [
        ("Architecture Layer", "Technologies Selected", "Design Rationale & Purpose"),
        ("Frontend Framework", "Next.js 14 (App Router) + React 18", "Server-side rendering, fast routing, built-in API proxying without CORS overhead."),
        ("Language & Types", "TypeScript 5.7 (Strict Mode)", "End-to-end type safety for geometric bounding boxes, schemas, and pipeline events."),
        ("UI Components & Styling", "Tailwind CSS 3.4 + Radix UI Primitives", "Fully accessible custom design system with dark/light themes and custom sliders."),
        ("PDF Client Rendering", "PDF.js 4.8.69", "Dual-canvas synchronized PDF rendering with coordinate mapping and zooming."),
        ("Backend Framework", "FastAPI 0.115 + Uvicorn 0.34", "High-concurrency asynchronous REST API with auto-generated OpenAPI documentation."),
        ("PDF & Vector Engine", "PyMuPDF (Fitz 1.25.2)", "C-level high-speed vector parsing, rawdict baseline extraction, and selective redaction."),
        ("Computer Vision & Imaging", "scikit-image 0.25 + Pillow 11.1", "Masked SSIM calculation, Sauvola adaptive thresholding, and median inpainting."),
        ("OCR & Image Processing", "pytesseract 0.3.13 + OpenCV / Hough", "Native-DPI rasterization, rotational deskewing, and word-level TSV confidence parsing."),
        ("Typography & Shaping", "fonttools 4.55, arabic-reshaper, python-bidi", "Font metric parsing, x-height normalization, RTL reversal, and Arabic ligature joining."),
        ("Database & Persistence", "SQLAlchemy 2.0 (SQLite / PostgreSQL)", "Immutable version records, segment storage, translation memory, and audit logging.")
    ]

    for row_idx, row_values in enumerate(tech_data):
        for col_idx, text in enumerate(row_values):
            tech_table.cell(row_idx, col_idx).text = text

    style_table(tech_table, [1.4, 2.1, 3.0], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    # SECTION 4
    add_h1("4. Deep Architectural Innovations")

    add_h2("4.1 Baseline-Accurate Vertical Registration")
    doc.add_paragraph(
        "Standard layout engines measure bounding boxes from the top edge, causing vertical text drift when font families change. "
        "LayoutLoom anchors placement directly to the typographic baseline:"
    )
    add_callout(
        "Top    = First_Baseline - (Ascender * Point_Size)\n"
        "Bottom = Max(bbox.y1, Last_Baseline - Descender * Point_Size) + 1.2pt\n\n"
        "This formula matches MuPDF's internal ascender coordinate system, eliminating vertical line drift across translations.",
        "VERTICAL REGISTRATION FORMULA"
    )

    add_h2("4.2 The 7-Rung Typographic Concession Ladder")
    doc.add_paragraph(
        "When translated text exceeds the original bounding box, LayoutLoom traverses a 7-rung concession ladder. "
        "The first rung that measurably fits wins:"
    )

    ladder_table = doc.add_table(rows=8, cols=4)
    ladder_data = [
        ("Rung", "Typographic Concession", "Boundary Constraints", "Penalty Weight"),
        ("0", "Original Font Size & Leading", "Zero modifications", "0"),
        ("1", "Tighten Line Leading", "Tighten to 1.00x, floor at 0.95x", "1"),
        ("2", "Tighten Letter-Tracking", "-2% em tracking (CSS; skipped for CJK)", "2"),
        ("3", "Reduce Font Point Size", "2% decrement steps (Floor: 0.82x; 0.75x in table cells)", "3"),
        ("4", "Downward Box Growth", "Expand into verified-empty space (Max +25% height)", "5"),
        ("5", "Outer Margin Growth", "Expand into page margin (Max 50% width; not in tables)", "8"),
        ("6", "Honest Overflow Recording", "Log exact point deficit as OVERFLOW issue (Never clipped)", "13")
    ]
    for r_idx, row in enumerate(ladder_data):
        for c_idx, val in enumerate(row):
            ladder_table.cell(r_idx, c_idx).text = val

    style_table(ladder_table, [0.8, 2.5, 2.4, 0.8], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

    add_h2("4.3 Masked SSIM Graphics Verification Oracle")
    doc.add_paragraph(
        "Standard SSIM fails on translated documents because changed words reduce the similarity score. "
        "LayoutLoom introduces Masked SSIM, which generates a binary mask covering all translatable text bboxes and computes "
        "Structural Similarity strictly over artwork, vector ruling lines, logos, and stamps:"
    )
    add_callout(
        "SSIM_masked = SSIM(Render_orig * Mask, Render_trans * Mask) >= 0.98\n\n"
        "Composite Layout Preservation Score = 100 * [ 0.40 * SSIM_masked + 0.25 * (1 - Overflow_Rate) "
        "+ 0.20 * (1 - Concession_Rate) + 0.15 * Geometry_Integrity ]",
        "MATHEMATICAL QUALITY METRIC"
    )

    # SECTION 5
    add_h1("5. Production Invariants (I1 – I8)")
    doc.add_paragraph("LayoutLoom enforces 8 strict correctness invariants asserted directly in code:")

    inv_table = doc.add_table(rows=9, cols=3)
    inv_data = [
        ("Invariant", "Requirement Rule", "Enforcement & Verification Location"),
        ("I1", "Output page count == Input page count", "Rebuilder._assert_geometry (Hard exception on mismatch)"),
        ("I2", "Per-page canvas dimensions == Input dimensions", "Rebuilder._assert_geometry (Bounding box coordinate check)"),
        ("I3", "Images and vector graphics remain untouched", "PDF_REDACT_IMAGE_NONE + Masked SSIM >= 0.98"),
        ("I4", "No silent text clipping or disappearance", "Scratch canvas fitting oracle logs Rung 6 overflow"),
        ("I5", "Zero placement overlap > 2% between text boxes", "Neighbor-aware rectangle trimming & per-page sweep"),
        ("I6", "Protected spans (URLs, codes, numbers) remain verbatim", "Token validation [Pn] with regex fallback guard"),
        ("I7", "Zero-dependency offline operation", "Pre-vendored Noto fonts, local OCR, and offline mock provider"),
        ("I8", "All quality scores are mathematically measured", "ValidationResult derivation functions (Zero hardcoded metrics)")
    ]
    for r_idx, row in enumerate(inv_data):
        for c_idx, val in enumerate(row):
            inv_table.cell(r_idx, c_idx).text = val

    style_table(inv_table, [1.0, 2.8, 2.7], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])

    # SECTION 6
    add_h1("6. Empirical Benchmark Results")
    doc.add_paragraph("Performance measured across the four standard reference documents:")

    bench_table = doc.add_table(rows=5, cols=8)
    bench_data = [
        ("Sample Document", "Target", "SSIM", "Coverage", "Overflow", "Overlaps", "Geom", "Score"),
        ("Govt Notice (Hindi, Vector Seal)", "en", "1.000", "1.00", "2", "0", "Pass", "89.1"),
        ("Research Paper (2-Col, Table)", "hi", "1.000", "1.00", "0", "0", "Pass", "98.6"),
        ("Technical Report (German, 3 Tables)", "en", "1.000", "1.00", "0", "0", "Pass", "98.2"),
        ("Scanned Invoice (200 DPI Scan)", "de", "1.000", "1.00", "5", "0", "Pass", "92.6")
    ]
    for r_idx, row in enumerate(bench_data):
        for c_idx, val in enumerate(row):
            bench_table.cell(r_idx, c_idx).text = val

    style_table(bench_table, [2.2, 0.6, 0.6, 0.7, 0.7, 0.7, 0.5, 0.5], 
                [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, 
                 WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, 
                 WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER])

    # SECTION 7
    add_h1("7. Presentation Q&A Defense & Strategy")
    doc.add_paragraph("Key answers for anticipated questions during technical reviews or pitch competitions:")

    add_bullet("Q: Why not use an LLM with vision to generate the translated PDF directly?", 
               "Vision LLMs cannot generate precise PDF vector commands, exact font baselines, or maintain 150+ DPI scan resolutions. LayoutLoom uses AI strictly for semantic translation while relying on deterministic C-level geometry engines for layout reconstruction.")
    add_bullet("Q: How does LayoutLoom prevent text overlapping with images?", 
               "Every non-text element (images, paths, vectors) is registered as an occupied bounding zone. Bounding box expansion in Rung 4/5 performs a neighbor collision sweep; if a collision is detected, the expansion reverts and proceeds to Rung 6.")
    add_bullet("Q: How do you handle non-standard fonts?", 
               "LayoutLoom parses PostScript font descriptors (family, weight, serif vs sans) and selects a matching pre-vendored Noto face, applying optical x-height normalization to ensure identical visual weight.")

    doc.save(output_path)
    print(f"Successfully generated styled Word document at: {output_path}")

if __name__ == "__main__":
    create_styled_document("/Users/sagarkumar_07/Desktop/PrototypeSIH/LayoutLoom_Pitch_Deck_and_Architecture_Guide.docx")
